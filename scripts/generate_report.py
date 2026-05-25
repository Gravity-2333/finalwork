from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "2300320125-杨翰飞-软件工程-23003201班-生产实习报告-模版.docx"
OUTPUT = ROOT / "2300320125-杨翰飞-软件工程-23003201班-生产实习报告.docx"
SCREENSHOTS = [
    ("人脸核验、Provider 配置与学习看板", ROOT / "screenshots" / "01-login-provider-dashboard.png"),
    ("知识库上传、大纲生成与章节工作区", ROOT / "screenshots" / "02-knowledge-outline-workspace.png"),
    ("章节学习内容生成结果", ROOT / "screenshots" / "03-chapter-content.png"),
    ("在线测验与错题归档", ROOT / "screenshots" / "04-quiz-wrong-answers.png"),
    ("语音交互控制入口", ROOT / "screenshots" / "05-voice-control.png"),
]


FILL_TEXT = {
    "purpose": (
        "本次生产实习项目选择开发“AI 学习助手”系统，目标是综合运用前后端开发、"
        "LangChain 文档处理、LangGraph 工作流编排和可切换大模型调用能力，完成一个能够服务课程学习的智能应用。"
        "系统围绕课程资料上传、知识库构建、学习大纲生成、章节内容学习、在线测验、错题复盘、人脸核验和语音交互展开。"
    ),
    "task": (
        "项目任务包括：搭建 FastAPI 后端与 Vue 3 前端；使用 LangChain 完成 txt、md、docx、pdf 学习资料加载与切割；"
        "使用 SQLite 保存知识库、章节进度、测验和错题；使用 LangGraph 编排检索、提示词组织和生成流程；"
        "实现 local_ollama、cloud_ollama、deepseek、mock 四类 Provider；将 LangSmith 设计为可选追踪功能；"
        "最终完成运行截图和实习报告整理。"
    ),
    "reflection": (
        "通过本项目，我对 AI 应用工程化流程有了更完整的认识。相比单独调用模型接口，真正可用的学习助手需要资料管理、"
        "状态持久化、错误降级、界面引导和证据化交付等多方面配合。Provider 架构让我理解到模型调用不应绑定单一服务，"
        "Mock fallback 也保证了演示和验收环境的稳定性。后续可继续扩展向量检索、真实人脸特征比对和更细粒度的学习分析。"
    ),
}


def replace_placeholder(document: Document, text: str) -> int:
    replaced = 0
    for paragraph in document.paragraphs:
        if "同学填写" in paragraph.text:
            paragraph.clear()
            paragraph.add_run(text)
            replaced += 1
            if replaced == 1:
                return replaced
    return replaced


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


def main() -> None:
    missing = [str(path) for _, path in SCREENSHOTS if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少截图文件：" + "；".join(missing))

    document = Document(TEMPLATE)
    replacements = [
        FILL_TEXT["purpose"],
        FILL_TEXT["task"],
        "本次实习单位填写为软通动力，实习岗位方向为 AI 应用开发与软件工程实践。",
        FILL_TEXT["reflection"],
    ]
    for text in replacements:
        replace_placeholder(document, text)

    document.add_page_break()
    add_heading(document, "AI 学习助手系统运行结果与截图说明")
    document.add_paragraph(
        "本系统位于 finalwork 目录，采用 FastAPI + Vue 3 前后端分离架构。"
        "后端使用 LangChain 完成课程资料加载和文本切割，使用 LangGraph 编排知识库检索、课程大纲、章节内容、测验生成和错题归档流程。"
        "模型调用支持 local_ollama、cloud_ollama、deepseek 和 mock 四种 Provider；缺少环境变量或接口不可用时，系统会给出友好提示并自动使用 Mock 演示。"
        "LangSmith 为可选追踪能力，默认关闭，没有 LANGSMITH_API_KEY 时不影响主流程。"
    )
    document.add_paragraph("关键运行截图如下：")

    for title, path in SCREENSHOTS:
        add_heading(document, title)
        document.add_picture(str(path), width=Inches(6.2))
        document.add_paragraph(f"截图说明：{title}，证明对应功能已经在系统中实现并可运行。")

    document.save(OUTPUT)
    print(f"Report generated: {OUTPUT}")


if __name__ == "__main__":
    main()

