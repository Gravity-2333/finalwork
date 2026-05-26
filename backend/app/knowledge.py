from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from .database import UPLOAD_DIR, connect


ALLOWED_SUFFIXES = {".txt", ".md", ".docx", ".pdf"}


def save_upload(file_name: str, content: bytes) -> Path:
    suffix = Path(file_name).suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise ValueError("仅支持 txt、md、docx、pdf 格式资料。")
    safe_name = re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", Path(file_name).name)
    target = UPLOAD_DIR / safe_name
    if target.exists():
        if target.read_bytes() == content:
            return target
        target = _next_available_path(target)
    target.write_bytes(content)
    return target


def build_knowledge(file_path: Path) -> dict:
    existing = _existing_document(file_path.name)
    if existing:
        return existing
    docs = _load_documents(file_path)
    splitter = RecursiveCharacterTextSplitter(chunk_size=650, chunk_overlap=90)
    chunks = splitter.split_documents(docs)
    summary = _summarize_chunks(chunks)
    with connect() as conn:
        cursor = conn.execute(
            "INSERT INTO documents(filename, file_path, chunk_count, summary) VALUES (?, ?, ?, ?)",
            (file_path.name, str(file_path), len(chunks), summary),
        )
        document_id = cursor.lastrowid
        conn.executemany(
            "INSERT INTO chunks(document_id, content, metadata) VALUES (?, ?, ?)",
            [
                (
                    document_id,
                    chunk.page_content,
                    json.dumps(chunk.metadata, ensure_ascii=False),
                )
                for chunk in chunks
            ],
        )
    return {
        "id": document_id,
        "filename": file_path.name,
        "chunk_count": len(chunks),
        "summary": summary,
    }


def list_documents() -> list[dict]:
    with connect() as conn:
        rows = conn.execute("SELECT * FROM documents ORDER BY id DESC").fetchall()
    documents = []
    for row in rows:
        document = dict(row)
        document["summary"] = clean_preview_text(document.get("summary", ""))
        documents.append(document)
    return documents


def delete_document(document_id: int) -> dict:
    with connect() as conn:
        row = conn.execute("SELECT * FROM documents WHERE id=?", (document_id,)).fetchone()
        if not row:
            raise ValueError("资料不存在或已被删除。")
        document = dict(row)
        conn.execute("DELETE FROM chunks WHERE document_id=?", (document_id,))
        conn.execute("DELETE FROM documents WHERE id=?", (document_id,))
        _clear_learning_state(conn)
    _unlink_if_unused(document["file_path"], document["filename"])
    return document


def clear_documents() -> int:
    with connect() as conn:
        rows = conn.execute("SELECT * FROM documents").fetchall()
        documents = [dict(row) for row in rows]
        conn.execute("DELETE FROM chunks")
        conn.execute("DELETE FROM documents")
        _clear_learning_state(conn)
    for document in documents:
        _unlink_if_unused(document["file_path"], document["filename"], assume_database_cleared=True)
    return len(documents)


def has_documents() -> bool:
    with connect() as conn:
        row = conn.execute("SELECT COUNT(*) AS count FROM chunks").fetchone()
    return bool(row and row["count"])


def search_chunks(query: str, limit: int = 8) -> list[dict]:
    terms = [item for item in re.split(r"\W+", query.lower()) if item]
    with connect() as conn:
        rows = conn.execute("SELECT id, document_id, content, metadata FROM chunks").fetchall()
    ranked = []
    for row in rows:
        content = row["content"]
        lowered = content.lower()
        score = sum(lowered.count(term) for term in terms) if terms else 0
        score += min(len(content), 650) / 10000
        ranked.append((score, dict(row)))
    ranked.sort(key=lambda item: item[0], reverse=True)
    return [item for score, item in ranked[:limit] if score > 0]


def all_context(limit: int = 12) -> str:
    with connect() as conn:
        rows = conn.execute("SELECT content FROM chunks ORDER BY id LIMIT ?", (limit,)).fetchall()
    return "\n\n".join(row["content"] for row in rows)


def knowledge_stats() -> dict:
    with connect() as conn:
        row = conn.execute("SELECT COUNT(*) AS chunks, COALESCE(SUM(LENGTH(content)), 0) AS chars FROM chunks").fetchone()
    return {"total_chunks": int(row["chunks"] or 0), "total_chars": int(row["chars"] or 0)}


def outline_context(limit: int = 24) -> str:
    keywords = ("目录", "第 1 章", "第1章", "第一章", "Part", "第一部分", "第二部分", "本章包括", "本章小结", "学习路线图", "关于本书")
    with connect() as conn:
        rows = conn.execute("SELECT id, content FROM chunks ORDER BY id").fetchall()
    scored = []
    for index, row in enumerate(rows):
        content = row["content"]
        score = 0
        score += max(0, 16 - index) * 0.35
        score += sum(5 for keyword in keywords if keyword in content)
        score += len(re.findall(r"(第\s*[一二三四五六七八九十0-9]+\s*[章节]|Chapter\s+\d+|Part\s+\d+)", content, re.IGNORECASE)) * 3
        scored.append((score, row["id"], content))
    scored.sort(key=lambda item: (-item[0], item[1]))
    selected = [content for score, _, content in scored[:limit] if score > 0]
    return "\n\n".join(selected) or all_context(limit)


def chapter_context(title: str, objective: str, limit: int = 14) -> str:
    chunks = search_chunks(f"{title} {objective}", limit)
    return "\n\n".join(item["content"] for item in chunks)


def document_outline_chapters() -> list[dict]:
    with connect() as conn:
        rows = conn.execute("SELECT file_path FROM documents ORDER BY id").fetchall()
    chapters = []
    for row in rows:
        path = Path(row["file_path"])
        if path.suffix.lower() != ".pdf" or not path.exists():
            continue
        chapters.extend(_pdf_outline_chapters(path))
    chapters = _dedupe_outline_chapters(chapters)
    if len(chapters) >= 4:
        return chapters
    return _text_outline_chapters()


def _pdf_outline_chapters(path: Path) -> list[dict]:
    try:
        reader = PdfReader(str(path))
        raw_outline = reader.outline
    except Exception:
        return []
    chapters = []

    def walk(items, depth: int = 0) -> None:
        for item in items:
            if isinstance(item, list):
                walk(item, depth + 1)
                continue
            title = getattr(item, "title", str(item)).replace("\r", "").strip()
            if depth <= 1 and re.match(r"^第\s*[一二三四五六七八九十0-9]+\s*章", title):
                clean_title = re.sub(r"\s+", " ", title)
                clean_title = re.sub(r"^第\s*([一二三四五六七八九十0-9]+)\s*章\s*", r"第\1章 ", clean_title).strip()
                chapters.append({"title": clean_title[:40], "objective": _objective_for_outline_title(clean_title)})

    walk(raw_outline)
    return chapters


def _objective_for_outline_title(title: str) -> str:
    topic = re.sub(r"^第\s*[一二三四五六七八九十0-9]+\s*章\s*", "", title).strip() or title
    return f"系统学习“{topic}”相关核心概念、方法流程、典型应用和易错点，能够完成本章测验与复盘。"


def _dedupe_outline_chapters(chapters: list[dict]) -> list[dict]:
    result = []
    seen = set()
    for chapter in chapters:
        key = chapter["title"]
        if key in seen:
            continue
        seen.add(key)
        result.append(chapter)
    return result


def _text_outline_chapters() -> list[dict]:
    with connect() as conn:
        rows = conn.execute("SELECT content FROM chunks ORDER BY id").fetchall()
    text = "\n".join(row["content"] for row in rows)
    candidates = []
    patterns = [
        r"(?m)^\s*(第\s*[一二三四五六七八九十0-9]+\s*章\s+[^\n]{2,40})\s*$",
        r"(?m)^\s*(Chapter\s+\d+[\s:：.-]+[^\n]{2,50})\s*$",
    ]
    for pattern in patterns:
        candidates.extend(match.group(1).strip() for match in re.finditer(pattern, text, flags=re.IGNORECASE))
    chapters = []
    seen_numbers = set()
    for title in candidates:
        normalized = _clean_outline_title(title)
        number_match = re.search(r"(?:第\s*([一二三四五六七八九十0-9]+)\s*章|Chapter\s+(\d+))", normalized, flags=re.IGNORECASE)
        number = (number_match.group(1) or number_match.group(2)) if number_match else normalized
        if number in seen_numbers:
            continue
        seen_numbers.add(number)
        chapters.append({"title": normalized[:40], "objective": _objective_for_outline_title(normalized)})
    return _dedupe_outline_chapters(chapters)


def _clean_outline_title(title: str) -> str:
    cleaned = re.sub(r"\s+", " ", title).strip()
    cleaned = re.sub(r"[.·。．]{3,}\s*\d*\s*$", "", cleaned)
    cleaned = re.sub(r"\s+\d{1,4}\s*$", "", cleaned)
    cleaned = re.sub(r"^第\s*([一二三四五六七八九十0-9]+)\s*章\s*", r"第\1章 ", cleaned)
    return cleaned.strip(" .-—")


def _existing_document(filename: str) -> dict | None:
    with connect() as conn:
        row = conn.execute("SELECT * FROM documents WHERE filename=? ORDER BY id DESC LIMIT 1", (filename,)).fetchone()
    return dict(row) if row else None


def _next_available_path(target: Path) -> Path:
    for index in range(1, 1000):
        candidate = target.with_name(f"{target.stem}_{index}{target.suffix}")
        if not candidate.exists():
            return candidate
    raise ValueError("同名资料过多，请重命名文件后再上传。")


def _load_documents(file_path: Path) -> list[Document]:
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _clean_documents(TextLoader(str(file_path), encoding="utf-8").load())
    if suffix == ".pdf":
        return _clean_documents(PyPDFLoader(str(file_path)).load())
    if suffix == ".docx":
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return _clean_documents([Document(page_content=text, metadata={"source": str(file_path)})])
    raise ValueError("不支持的资料格式。")


def _clean_documents(documents: list[Document]) -> list[Document]:
    cleaned = []
    for doc in documents:
        text = doc.page_content.encode("utf-8", "ignore").decode("utf-8")
        cleaned.append(Document(page_content=text, metadata=doc.metadata))
    return cleaned


def _summarize_chunks(chunks: list[Document]) -> str:
    text = "\n".join(chunk.page_content for chunk in chunks[:3])
    return clean_preview_text(text) or "资料已入库，部分页面可能因扫描件或字体编码导致预览不可读。"


def clean_preview_text(text: str) -> str:
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text or "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = re.sub(r"[^A-Za-z0-9_\u4e00-\u9fff，。！？、；：,.!?;:()（）《》“”\"'\-\s/%]+", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    first_cjk = re.search(r"[\u4e00-\u9fff]", cleaned)
    if first_cjk and first_cjk.start() > 16 and is_fragmented_prefix(cleaned[: first_cjk.start()]):
        cleaned = cleaned[first_cjk.start() :].strip()
    preview = cleaned[:240].strip()
    if is_probably_garbled(preview):
        return "资料已入库，部分 PDF 页面可能因扫描件或字体编码导致预览不可读，但系统会尽量使用可提取文本生成学习内容。"
    return preview[:220]


def is_probably_garbled(text: str) -> bool:
    if not text:
        return True
    if text.count("�") >= 3:
        return True
    readable = re.findall(r"[A-Za-z0-9_\u4e00-\u9fff，。！？、；：,.!?;:()（）《》“”\"'\-\s/%]", text)
    readable_ratio = len(readable) / max(len(text), 1)
    cjk_or_word = re.findall(r"[A-Za-z0-9_\u4e00-\u9fff]", text)
    content_ratio = len(cjk_or_word) / max(len(text), 1)
    abnormal_runs = re.search(r"[^A-Za-z0-9_\u4e00-\u9fff\s]{8,}", text)
    tokens = re.findall(r"[A-Za-z0-9_\u4e00-\u9fff]+", text)
    short_tokens = [token for token in tokens if len(token) <= 2]
    cjk_chars = re.findall(r"[\u4e00-\u9fff]", text)
    short_token_ratio = len(short_tokens) / max(len(tokens), 1)
    fragmented_latin = len(tokens) >= 12 and short_token_ratio > 0.6 and len(cjk_chars) < 12
    return readable_ratio < 0.72 or content_ratio < 0.25 or bool(abnormal_runs) or fragmented_latin


def is_fragmented_prefix(text: str) -> bool:
    tokens = re.findall(r"[A-Za-z0-9_]+", text)
    if len(tokens) < 6:
        return False
    short_tokens = [token for token in tokens if len(token) <= 2]
    return len(short_tokens) / max(len(tokens), 1) > 0.55


def _clear_learning_state(conn) -> None:
    conn.execute("DELETE FROM wrong_answers")
    conn.execute("DELETE FROM quizzes")
    conn.execute("DELETE FROM chapters")


def _unlink_if_unused(file_path: str, filename: str, assume_database_cleared: bool = False) -> None:
    target = Path(file_path)
    if not target.exists() or target.parent != UPLOAD_DIR:
        return
    if not assume_database_cleared and _existing_document(filename):
        return
    target.unlink(missing_ok=True)
